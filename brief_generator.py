#!/usr/bin/env python3
"""
Pre-Call Research Brief Generator
Powered by Claude AI — Metabase GTM Toolkit

SETUP (one time):
  pip install -r requirements.txt
  export ANTHROPIC_API_KEY=your_key_here   ← get it at console.anthropic.com

USAGE:
  python brief_generator.py
"""

import os
import sys
import re
import anthropic
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional: richer news context. Install with: pip install duckduckgo-search
try:
    from duckduckgo_search import DDGS
    HAS_SEARCH = True
except ImportError:
    HAS_SEARCH = False


def fetch_website(url):
    """Pull visible text from a company's homepage."""
    if not url:
        return ""
    try:
        if not url.startswith("http"):
            url = "https://" + url
        headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
        r = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        # Strip nav, footer, scripts — we just want the content
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)
        return text[:4000]  # Enough context without blowing up the prompt
    except Exception as e:
        return f"(Could not fetch website: {e})"


def fetch_news(company_name):
    """Pull recent news headlines about the company (requires duckduckgo-search)."""
    if not HAS_SEARCH:
        return ""
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.news(company_name, max_results=5):
                results.append(f"- {r['title']} ({r.get('date', 'recent')}): {r['body'][:200]}")
        return "\n".join(results) if results else ""
    except Exception:
        return ""


def generate_brief(company_name, website_url, contact_name, contact_title):
    """Call Claude to generate the research brief."""
    website_content = fetch_website(website_url)
    news_content    = fetch_news(company_name)

    context_parts = []
    if website_content:
        context_parts.append(f"Website content:\n{website_content}")
    if news_content:
        context_parts.append(f"Recent news:\n{news_content}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are helping an Account Executive at Metabase prepare for a discovery call.

Metabase is the leading open-source business intelligence and embedded analytics platform. It lets anyone
in an organization explore data and build dashboards without writing SQL — and lets SaaS companies embed
beautiful, interactive analytics directly inside their own product. Metabase runs on every major database
and data warehouse, and is used by over 60,000 organizations worldwide.

Metabase has two core use cases:
1. Self-serve internal BI: Non-technical teams (product, marketing, ops, finance) can answer their own
   data questions without filing tickets to the data team.
2. Embedded analytics: SaaS/product companies embed Metabase dashboards inside their own product to give
   customers data visibility — without building a custom analytics layer from scratch.

---
PROSPECT INFO:
Company:       {company_name}
Contact:       {contact_name or "Unknown"} — {contact_title or "Unknown title"}
---
CONTEXT:
{context}

---
Generate a tight, scannable pre-call research brief. Format it exactly like this:

## {company_name} — Pre-Call Brief

### What They Do
[2–3 sentences on the company's core business and scale]

### Likely Data Stack
[Based on available signals, what databases, warehouses, or BI tools are they likely running? Be specific if you can infer it.]

### Analytics & BI Pain Points
[What data/analytics challenges might a company like this face? Think: analyst bottlenecks, dashboard sprawl, embedded reporting needs, expensive legacy BI tools, non-technical users blocked from data. Root cause, not surface symptoms.]

### Metabase Fit
[How could Metabase specifically solve their problems — be concrete. Is this a self-serve BI play, an embedded analytics play, or both? What value does Metabase deliver in their context?]

### Tailored Discovery Questions
1. [Question]
2. [Question]
3. [Question]
4. [Question]
5. [Question]

### One-Line Opener
[A compelling, personalized opening line for the first 30 seconds of the call — sounds human, not scripted]

Keep the whole brief under 450 words. Make it something the AE can read in 90 seconds right before the call."""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_cold_emails(company_name, website_url, contact_name, contact_title, custom_notes):
    """Generate 3 cold email variations for a Metabase prospect."""
    website_content = fetch_website(website_url)
    news_content    = fetch_news(company_name)

    context_parts = []
    if website_content:
        context_parts.append(f"Website content:\n{website_content}")
    if news_content:
        context_parts.append(f"Recent news:\n{news_content}")
    if custom_notes:
        context_parts.append(f"Additional context from the rep:\n{custom_notes}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are helping an Account Executive at Metabase write cold outbound emails.

Metabase is the open-source BI and embedded analytics platform that lets every team explore data without
SQL, and lets SaaS companies embed analytics inside their product without building it from scratch.
It replaces expensive, complex BI tools (Tableau, Looker) and eliminates the analyst bottleneck that
slows most data-driven companies down.

---
PROSPECT INFO:
Company:       {company_name}
Contact:       {contact_name or "Unknown"} — {contact_title or "Unknown title"}
---
CONTEXT:
{context}

---
Write 3 cold email variations. Each needs a subject line and a short body.

Format exactly like this:

---
**Variation 1: Direct**
**Subject:** [subject line]

[Body]

---
**Variation 2: Insight-Led**
**Subject:** [subject line]

[Body]

---
**Variation 3: Question-Based**
**Subject:** [subject line]

[Body]

---
Rules — follow every one of these:
- Each email must be under 100 words
- Use short paragraphs of 1 to 2 lines each
- Each email follows one simple structure: one observation, one specific pain (analyst bottleneck, dashboard sprawl, embedded reporting, expensive BI tool, or similar), one outcome, one ask
- Never stack multiple ideas or features in the same email
- Use concrete, specific language. No "purpose-built", "robust", "seamlessly", or generic sales phrases
- Write like a peer who has seen this problem before at their scale, not like a salesperson pitching a product
- No dashes used as punctuation. No overly polished transitions
- Sign off with: [Your name] at Metabase
- One CTA per email, kept light: "curious if you've run into this?" or "worth a quick call?" — never "book a demo"
- The tone should feel like it was written in 30 seconds, not edited for 30 minutes
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_crm_summary(call_notes, company_name, contact_name, contact_title, call_date):
    """Generate a CRM-ready call summary with standard fields and MEDDIC scoring."""

    prompt = f"""You are helping an Account Executive at Metabase log a sales call in their CRM.

Metabase is the open-source BI and embedded analytics platform — used by 60,000+ organizations to give
every team self-serve data access and to let SaaS companies embed analytics inside their own products.

---
CALL INFO:
Company:      {company_name or "Unknown"}
Contact:      {contact_name or "Unknown"} — {contact_title or "Unknown title"}
Date:         {call_date or "Unknown"}
---
CALL NOTES:
{call_notes}

---
Generate a structured CRM summary in two parts.

PART 1 — CALL SUMMARY (for the CRM notes field):

### Overview
[2–3 sentences capturing what was discussed and the outcome of the call]

### Pain Points
[Bullet list of specific pains the prospect mentioned — e.g. analyst bottleneck, embedded reporting needs, Tableau/Looker cost, non-technical teams blocked from data]

### Objections
[Bullet list of objections raised, or "None raised" if none]

### Stakeholders Mentioned
[Name — Title — Role in deal, one per line. "None mentioned" if none]

### Next Steps
[Numbered list of agreed next steps with owners]

### Deal Stage
[Recommended CRM stage and one sentence of reasoning]

---

PART 2 — MEDDIC SCORECARD:

**Metrics:** [What quantifiable outcomes did they mention? Hours saved, dashboards replaced, analyst headcount, cost of current BI tool, time-to-insight?]
**Economic Buyer:** [Who controls the budget? Identified or unknown?]
**Decision Criteria:** [What will they use to evaluate solutions? Open source, ease of use, embedding capability, cost vs Tableau/Looker?]
**Decision Process:** [How will they decide? Timeline? Who else is involved?]
**Identify Pain:** [What is the core business pain driving this evaluation — self-serve BI bottleneck, embedded analytics need, or BI tool cost/complexity?]
**Champion:** [Is there an internal advocate? Who?]

---
Rules:
- Be specific, not generic. Use exact words or numbers from the call notes where possible.
- If something wasn't mentioned, write "Not discussed" rather than guessing.
- Keep the tone professional and factual — this is going into a CRM, not a pitch deck.
- Total length should be under 400 words.
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_icp_score(company_name, website_url):
    """Research a company and score it against Metabase's ICP (1-4)."""
    website_content = fetch_website(website_url)
    news_content    = fetch_news(company_name)

    context_parts = []
    if website_content:
        context_parts.append(f"Website content:\n{website_content}")
    if news_content:
        context_parts.append(f"Recent news:\n{news_content}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are a sales researcher helping Metabase's GTM team decide whether to pursue a prospect.

Metabase's Ideal Customer Profile (ICP):

SELF-SERVE BI USE CASE:
- Companies with a data team (even 1–2 analysts) that is a bottleneck for other teams
- Non-technical stakeholders (product, marketing, ops, finance) who need data but can't get it fast enough
- Companies using spreadsheets, homegrown dashboards, or legacy tools (Tableau, Looker) that are too slow or expensive
- Industries: SaaS, fintech, e-commerce, marketplace, healthcare tech, media, logistics
- Stage: Series A and beyond, or any established company with operational data they want to act on
- Pain signals: long BI request queues, engineers spending time building internal dashboards, high Tableau/Looker bills, scattered reporting across tools

EMBEDDED ANALYTICS USE CASE:
- SaaS or product companies that need to show analytics/reporting to their own customers
- Engineering teams currently building custom analytics dashboards from scratch
- Companies where "customer-facing reporting" is on the product roadmap but deprioritized due to build cost
- Pain signals: engineers maintaining custom chart code, customers asking for better reporting, product teams scoping a BI-within-app feature

NOT a good fit:
- Companies with no structured data or very minimal data needs
- Pure front-end or content companies with no operational/product data
- Companies fully locked into a BI platform with no budget or appetite to change
- Early-stage startups with no data infrastructure or analytics need yet
- Companies that only need a data warehouse or ETL tool (not a BI/visualization layer)

---
COMPANY: {company_name}
WEBSITE: {website_url or "Not provided"}
---
CONTEXT:
{context}

---
Research this company and produce an ICP scorecard. Format it exactly like this:

### ICP Score: [1, 2, 3, or 4] / 4

**[One-line verdict — e.g. "Strong fit — pursue now" or "Poor fit — deprioritize"]**

---

### Why This Score

[2–3 sentences explaining the reasoning. Be specific about what signals you found or didn't find. Note whether this is a self-serve BI fit, embedded analytics fit, or both.]

---

### ICP Signal Breakdown

- **Internal BI Need:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Embedded Analytics Need:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Data Team / Analyst Bottleneck:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Industry Fit:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Company Stage & Scale:** [Strong / Moderate / Weak / Unknown] — [one line explanation]

---

### Recommended Action

[One of these four, based on the score:]
- **Score 4:** Prioritize immediately. Add to active pipeline and reach out this week.
- **Score 3:** Worth pursuing. Research further and add to outbound sequence.
- **Score 2:** Possible fit. Monitor and revisit when you have more information.
- **Score 1:** Deprioritize. Move on — better opportunities exist.

---

Scoring guide:
- 4 = 4–5 strong ICP signals — clear fit, high priority
- 3 = 3 strong signals or 4–5 moderate ones — good fit, worth pursuing
- 2 = 1–2 strong signals or mixed signals — possible fit, needs qualification
- 1 = Few or no ICP signals — poor fit, not worth time right now

Be honest. If a company doesn't fit, say so clearly. A bad lead wastes more time than no lead.
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def parse_email_variations(emails_text):
    """Parse AI email output into a list of (title, subject, body) tuples."""
    variations = []
    chunks = re.split(r'\n?---\n?', emails_text)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        lines = chunk.split('\n')
        title, subject, body_lines, found_subject = "", "", [], False

        for line in lines:
            clean = line.strip()
            if re.match(r'\*\*Variation \d+', clean):
                title = re.sub(r'\*\*', '', clean).strip()
            elif clean.startswith('**Subject:**'):
                subject = clean.replace('**Subject:**', '').strip()
                found_subject = True
            elif found_subject:
                body_lines.append(line)

        if title and subject:
            body = '\n'.join(body_lines).strip()
            body = re.sub(r'\n{3,}', '\n\n', body)
            variations.append((title, subject, body))

    return variations


def save_as_docx(brief_text, filepath):
    """Convert the markdown-style brief into a formatted Word document."""
    doc = Document()

    # Document-wide style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in brief_text.splitlines():
        line = line.strip()

        if line.startswith("## "):
            # Main title
            p = doc.add_heading(line[3:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x50, 0x9E, 0xE3)  # Metabase blue

        elif line.startswith("### "):
            # Section heading
            doc.add_heading(line[4:], level=2)

        elif re.match(r"^\*(.+)\*$", line):
            # *italic date line*
            p = doc.add_paragraph()
            run = p.add_run(re.match(r"^\*(.+)\*$", line).group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        elif re.match(r"^\d+\.", line):
            # Numbered list item — use plain paragraph to avoid double-numbering
            p = doc.add_paragraph(style="List Bullet")
            p.style = doc.styles["Normal"]
            p.add_run(line)

        elif line == "" or line == "---":
            # Blank line / divider → small spacer
            doc.add_paragraph("")

        else:
            doc.add_paragraph(line)

    doc.save(filepath)


def main():
    print("\n╔══════════════════════════════════════╗")
    print("║   Pre-Call Research Brief Generator  ║")
    print("║           Metabase GTM Toolkit        ║")
    print("╚══════════════════════════════════════╝\n")

    # Check for API key upfront
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("⚠️  ANTHROPIC_API_KEY not set.")
        print("   Get your key at: https://console.anthropic.com")
        print("   Then run:  export ANTHROPIC_API_KEY=your_key_here\n")
        sys.exit(1)

    # Collect inputs
    company_name  = input("Company name:   ").strip()
    if not company_name:
        print("Company name is required.")
        sys.exit(1)

    website_url   = input("Website URL:    ").strip()
    contact_name  = input("Contact name:   ").strip()
    contact_title = input("Contact title:  ").strip()

    # Status messages
    print(f"\n⏳ Researching {company_name}...")
    if HAS_SEARCH:
        print("   → Fetching website & recent news...")
    else:
        print("   → Fetching website content...")
    print("   → Generating brief with Claude...\n")

    # Generate
    try:
        brief = generate_brief(company_name, website_url, contact_name, contact_title)
    except anthropic.AuthenticationError:
        print("❌ Invalid API key. Double-check your ANTHROPIC_API_KEY.")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        sys.exit(1)

    # Print to terminal
    print("─" * 50)
    print(brief)
    print("─" * 50)

    # Save as a Word doc in a company folder on the Desktop
    desktop     = os.path.join(os.path.expanduser("~"), "Desktop")
    safe_name   = company_name.replace("/", "").strip()
    folder_path = os.path.join(desktop, safe_name)
    os.makedirs(folder_path, exist_ok=True)

    filename = os.path.join(folder_path, "brief.docx")
    save_as_docx(brief, filename)

    print(f"\n✅ Saved to Desktop → {safe_name} → brief.docx")
    print("   Open it in Word or drag it into Google Docs.\n")


if __name__ == "__main__":
    main()
